"""
文件操作模块 - 文件读写、Excel统计、Word生成、文件夹整理

为智能体提供完整的文件系统操作能力：
- 读取/写入文本文件
- 读取/统计 Excel 数据
- 生成 Word 文档
- 文件夹整理（按日期/类型分类）
- 搜索文件
"""

import os
import json
import csv
import shutil
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any

logger = logging.getLogger(__name__)


class FileOperator:
    """文件操作管理器。"""

    def __init__(self):
        self._op_counter = 0

    # ---------- 文本文件 ----------

    def read_text(self, file_path: str, encoding: str = "utf-8",
                  start_line: int = 0, end_line: int = -1) -> Dict[str, Any]:
        """读取文本文件内容。"""
        try:
            with open(file_path, "r", encoding=encoding) as f:
                lines = f.readlines()
            total_lines = len(lines)
            if end_line < 0 or end_line > total_lines:
                end_line = total_lines
            selected = lines[start_line:end_line]
            content = "".join(selected)
            return {
                "success": True,
                "content": content,
                "total_lines": total_lines,
                "read_lines": len(selected),
                "file_path": os.path.abspath(file_path),
            }
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    lines = f.readlines()
                content = "".join(lines[start_line:end_line if end_line > 0 else len(lines)])
                return {
                    "success": True,
                    "content": content,
                    "total_lines": len(lines),
                    "read_lines": len(lines),
                    "file_path": os.path.abspath(file_path),
                    "encoding_used": "gbk",
                }
            except Exception as e:
                return {"success": False, "error": f"读取失败: {e}"}
        except Exception as e:
            return {"success": False, "error": f"读取失败: {e}"}

    def write_text(self, file_path: str, content: str,
                   mode: str = "w", encoding: str = "utf-8") -> Dict[str, Any]:
        """写入文本文件。"""
        try:
            os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)
            with open(file_path, mode, encoding=encoding) as f:
                f.write(content)
            return {
                "success": True,
                "file_path": os.path.abspath(file_path),
                "bytes_written": len(content.encode(encoding)),
            }
        except Exception as e:
            return {"success": False, "error": f"写入失败: {e}"}

    def append_text(self, file_path: str, content: str) -> Dict[str, Any]:
        """追加写入文本文件。"""
        return self.write_text(file_path, content, mode="a")

    # ---------- Excel 操作 ----------

    def read_excel(self, file_path: str, sheet_name: Optional[str] = None,
                   sheet_index: int = 0, header: bool = True) -> Dict[str, Any]:
        """读取 Excel 文件。"""
        try:
            import openpyxl
        except ImportError:
            return {"success": False, "error": "openpyxl 未安装，请先 pip install openpyxl"}

        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            if sheet_name:
                ws = wb[sheet_name]
            else:
                ws = wb.worksheets[sheet_index]

            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(cell) if cell is not None else "" for cell in row])

            headers = rows[0] if header and rows else []
            data = rows[1:] if header and len(rows) > 1 else rows

            return {
                "success": True,
                "sheet_name": ws.title,
                "total_rows": ws.max_row,
                "total_cols": ws.max_column,
                "headers": headers,
                "data": data,
                "row_count": len(data),
                "file_path": os.path.abspath(file_path),
            }
        except Exception as e:
            return {"success": False, "error": f"Excel 读取失败: {e}"}

    def excel_statistics(self, file_path: str, sheet_name: Optional[str] = None) -> Dict[str, Any]:
        """对 Excel 文件进行统计分析。"""
        result = self.read_excel(file_path, sheet_name=sheet_name)
        if not result.get("success"):
            return result

        data = result.get("data", [])
        headers = result.get("headers", [])

        if not data:
            return {"success": True, "statistics": {"message": "文件为空"}}

        # 转换为字典格式便于分析
        records = []
        for row in data:
            record = {}
            for i, val in enumerate(row):
                key = headers[i] if i < len(headers) else f"列{i+1}"
                record[key] = val
            records.append(record)

        # 统计每列的数值
        col_stats = {}
        if headers:
            for col_idx, header in enumerate(headers):
                numeric_values = []
                non_empty = 0
                for row in data:
                    if col_idx < len(row) and row[col_idx]:
                        non_empty += 1
                        try:
                            val = float(row[col_idx])
                            numeric_values.append(val)
                        except (ValueError, TypeError):
                            pass

                stats = {
                    "非空单元格": non_empty,
                    "空单元格": len(data) - non_empty,
                }
                if numeric_values:
                    stats["数值统计"] = {
                        "数量": len(numeric_values),
                        "总和": round(sum(numeric_values), 2),
                        "平均": round(sum(numeric_values) / len(numeric_values), 2),
                        "最小": round(min(numeric_values), 2),
                        "最大": round(max(numeric_values), 2),
                    }
                col_stats[header] = stats

        return {
            "success": True,
            "file_path": os.path.abspath(file_path),
            "sheet_name": result.get("sheet_name", ""),
            "total_rows": result.get("total_rows", 0),
            "total_cols": result.get("total_cols", 0),
            "columns": col_stats,
            "数据预览": records[:5],
        }

    # ---------- Word 生成 ----------

    def create_word(self, file_path: str, title: str = "",
                    paragraphs: Optional[List[str]] = None,
                    table_data: Optional[List[List[str]]] = None) -> Dict[str, Any]:
        """生成 Word 文档。"""
        try:
            from docx import Document
            from docx.shared import Pt, Inches
        except ImportError:
            return {"success": False, "error": "python-docx 未安装，请先 pip install python-docx"}

        try:
            doc = Document()

            if title:
                heading = doc.add_heading(title, level=0)

            if paragraphs:
                for para_text in paragraphs:
                    p = doc.add_paragraph(para_text)

            if table_data and len(table_data) > 0:
                rows = len(table_data)
                cols = len(table_data[0]) if table_data else 0
                table = doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"
                for i, row_data in enumerate(table_data):
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            table.rows[i].cells[j].text = str(cell_text)

            os.makedirs(os.path.dirname(file_path) or ".", exist_ok=True)
            doc.save(file_path)

            return {
                "success": True,
                "file_path": os.path.abspath(file_path),
                "size_bytes": os.path.getsize(file_path),
            }
        except Exception as e:
            return {"success": False, "error": f"Word 生成失败: {e}"}

    # ---------- 文件夹操作 ----------

    def list_files(self, folder_path: str, pattern: str = "*",
                   recursive: bool = False, max_results: int = 200) -> Dict[str, Any]:
        """列出文件夹中的文件。"""
        try:
            import fnmatch
            result = []
            if recursive:
                for root, dirs, files in os.walk(folder_path):
                    for f in files:
                        if fnmatch.fnmatch(f.lower(), pattern.lower()):
                            full_path = os.path.join(root, f)
                            stat = os.stat(full_path)
                            result.append({
                                "name": f,
                                "path": full_path,
                                "size": stat.st_size,
                                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                            })
            else:
                for item in os.listdir(folder_path):
                    full_path = os.path.join(folder_path, item)
                    if os.path.isfile(full_path) and fnmatch.fnmatch(item.lower(), pattern.lower()):
                        stat = os.stat(full_path)
                        result.append({
                            "name": item,
                            "path": full_path,
                            "size": stat.st_size,
                            "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                        })

            result.sort(key=lambda x: x["modified"], reverse=True)
            result = result[:max_results]

            return {
                "success": True,
                "folder": os.path.abspath(folder_path),
                "file_count": len(result),
                "files": result,
            }
        except Exception as e:
            return {"success": False, "error": f"列出文件失败: {e}"}

    def search_files(self, folder_path: str, keyword: str,
                     recursive: bool = True) -> Dict[str, Any]:
        """按关键字搜索文件。"""
        try:
            result = []
            search_lower = keyword.lower()
            if recursive:
                for root, dirs, files in os.walk(folder_path):
                    for f in files:
                        if search_lower in f.lower():
                            full_path = os.path.join(root, f)
                            stat = os.stat(full_path)
                            result.append({
                                "name": f,
                                "path": full_path,
                                "size": stat.st_size,
                                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                            })
            else:
                for item in os.listdir(folder_path):
                    if search_lower in item.lower():
                        full_path = os.path.join(folder_path, item)
                        if os.path.isfile(full_path):
                            stat = os.stat(full_path)
                            result.append({
                                "name": item,
                                "path": full_path,
                                "size": stat.st_size,
                                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                            })

            return {
                "success": True,
                "folder": os.path.abspath(folder_path),
                "keyword": keyword,
                "match_count": len(result),
                "matches": result[:50],
            }
        except Exception as e:
            return {"success": False, "error": f"搜索失败: {e}"}

    def organize_by_date(self, folder_path: str, target_folder: Optional[str] = None,
                         recursive: bool = False) -> Dict[str, Any]:
        """按修改日期整理文件（移动到对应日期的子文件夹）。"""
        try:
            if target_folder is None:
                target_folder = os.path.join(folder_path, "_按日期整理")
            os.makedirs(target_folder, exist_ok=True)

            moved_count = 0
            skipped = []
            items = os.listdir(folder_path) if not recursive else []

            if recursive:
                for root, dirs, files in os.walk(folder_path):
                    for f in files:
                        full_path = os.path.join(root, f)
                        if not full_path.startswith(target_folder):
                            stat = os.stat(full_path)
                            date_folder = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m")
                            dest_dir = os.path.join(target_folder, date_folder)
                            os.makedirs(dest_dir, exist_ok=True)
                            dest_path = os.path.join(dest_dir, f)
                            if not os.path.exists(dest_path):
                                shutil.move(full_path, dest_path)
                                moved_count += 1
                            else:
                                skipped.append(f"{f} (已存在)")
            else:
                for item in items:
                    full_path = os.path.join(folder_path, item)
                    if not os.path.isfile(full_path) or item.startswith("_"):
                        continue
                    stat = os.stat(full_path)
                    date_folder = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m")
                    dest_dir = os.path.join(target_folder, date_folder)
                    os.makedirs(dest_dir, exist_ok=True)
                    dest_path = os.path.join(dest_dir, item)
                    if not os.path.exists(dest_path):
                        shutil.move(full_path, dest_path)
                        moved_count += 1
                    else:
                        skipped.append(f"{item} (已存在)")

            return {
                "success": True,
                "moved_count": moved_count,
                "skipped": skipped[:20],
                "target_folder": os.path.abspath(target_folder),
            }
        except Exception as e:
            return {"success": False, "error": f"整理失败: {e}"}

    def organize_by_type(self, folder_path: str, target_folder: Optional[str] = None) -> Dict[str, Any]:
        """按文件类型整理文件。"""
        try:
            if target_folder is None:
                target_folder = os.path.join(folder_path, "_按类型整理")
            os.makedirs(target_folder, exist_ok=True)

            type_map = {
                "图片": {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff"},
                "文档": {".doc", ".docx", ".pdf", ".txt", ".rtf", ".md"},
                "表格": {".xls", ".xlsx", ".csv"},
                "演示": {".ppt", ".pptx"},
                "音频": {".mp3", ".wav", ".flac", ".aac", ".ogg"},
                "视频": {".mp4", ".avi", ".mkv", ".mov", ".wmv", ".flv"},
                "压缩包": {".zip", ".rar", ".7z", ".tar", ".gz"},
                "代码": {".py", ".js", ".ts", ".java", ".cpp", ".c", ".html", ".css"},
            }

            def classify_ext(ext: str) -> str:
                ext_lower = ext.lower()
                for category, extensions in type_map.items():
                    if ext_lower in extensions:
                        return category
                return "其他"

            moved_count = 0
            skipped = []

            for item in os.listdir(folder_path):
                full_path = os.path.join(folder_path, item)
                if not os.path.isfile(full_path) or item.startswith("_"):
                    continue

                _, ext = os.path.splitext(item)
                category = classify_ext(ext)
                dest_dir = os.path.join(target_folder, category)
                os.makedirs(dest_dir, exist_ok=True)
                dest_path = os.path.join(dest_dir, item)

                if not os.path.exists(dest_path):
                    shutil.move(full_path, dest_path)
                    moved_count += 1
                else:
                    skipped.append(f"{item} (已存在)")

            return {
                "success": True,
                "moved_count": moved_count,
                "skipped": skipped[:20],
                "target_folder": os.path.abspath(target_folder),
            }
        except Exception as e:
            return {"success": False, "error": f"整理失败: {e}"}

    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """获取文件详细信息。"""
        try:
            stat = os.stat(file_path)
            _, ext = os.path.splitext(file_path)
            return {
                "success": True,
                "name": os.path.basename(file_path),
                "path": os.path.abspath(file_path),
                "extension": ext,
                "size_bytes": stat.st_size,
                "size_readable": self._format_size(stat.st_size),
                "created": datetime.fromtimestamp(stat.st_ctime).strftime("%Y-%m-%d %H:%M:%S"),
                "modified": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
                "accessed": datetime.fromtimestamp(stat.st_atime).strftime("%Y-%m-%d %H:%M:%S"),
                "is_file": os.path.isfile(file_path),
                "is_dir": os.path.isdir(file_path),
            }
        except Exception as e:
            return {"success": False, "error": f"获取信息失败: {e}"}

    @staticmethod
    def _format_size(size_bytes: int) -> str:
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        elif size_bytes < 1024 * 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.1f} MB"
        else:
            return f"{size_bytes / (1024 * 1024 * 1024):.1f} GB"

    def read_csv(self, file_path: str, delimiter: str = ",") -> Dict[str, Any]:
        """读取 CSV 文件。"""
        try:
            with open(file_path, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f, delimiter=delimiter)
                rows = list(reader)
            if not rows:
                return {"success": True, "headers": [], "data": [], "row_count": 0}
            headers = rows[0]
            data = rows[1:]
            return {
                "success": True,
                "headers": headers,
                "data": data,
                "row_count": len(data),
                "file_path": os.path.abspath(file_path),
            }
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    reader = csv.reader(f, delimiter=delimiter)
                    rows = list(reader)
                headers = rows[0] if rows else []
                data = rows[1:] if len(rows) > 1 else []
                return {
                    "success": True,
                    "headers": headers,
                    "data": data,
                    "row_count": len(data),
                    "file_path": os.path.abspath(file_path),
                    "encoding_used": "gbk",
                }
            except Exception as e:
                return {"success": False, "error": f"CSV 读取失败: {e}"}
        except Exception as e:
            return {"success": False, "error": f"CSV 读取失败: {e}"}
